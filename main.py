"""
Реализация маляра для .docx файлов в виде ящика с инструментами.
(пока только одна отвёртка и та крестовая)
"""

from copy import deepcopy
from typing import Generator, List, Tuple

from docx import Document
from docx.text.paragraph import Paragraph, Run

from utility import Color, Index

__all__ = ('color_text',)


def color_text(document: Document,
               text: str,
               color: str = 'red',
               first_only: bool = False
               ) -> None:
    """
    Функция для покраски частей текста в .docx
    без изменения структуры и стилей.
    Покраска происходит на месте, не забудьте сохранить документ в файл.

    :param document: Экземпляр документа, который красим.
    :param text: Строка текста, которую нужно покрасить.
    :param color: Цвет (из класса Color), в который хотим покрасить.
    :param first_only: Флаг для покраски только первого вхождения.
    """
    text = text.strip()
    for paragraph in document.paragraphs:
        if not __find_text(paragraph, text, strict=False):
            continue
        for run in __get_runs_to_color(
                paragraph, text, first_only=first_only):
            __color_run(run, color)


def __find_text(element: Run | Paragraph,
                text: str,
                strict: bool = True
                ) -> bool:
    if strict:
        return text == element.text.strip()
    return text in element.text


def __get_runs_to_color(paragraph: Paragraph,
                        text: str,
                        start: int = 0,
                        first_only: bool = False
                        ) -> List[Run]:
    runs_to_color = []
    for run, text_part in __find_text_in_runs(paragraph.runs[start:], text):
        if __find_text(run, text_part, strict=True):
            runs_to_color.append(run)
            if first_only:
                return runs_to_color
            continue
        if __find_text(run, text_part, strict=False):
            start = [run.text for run in paragraph.runs].index(run.text)
            runs_to_color.append(__reshape_run_with_text(
                paragraph, run, text_part))
            if first_only:
                return runs_to_color
            runs_to_color += __get_runs_to_color(paragraph, text, start)
            break
    return runs_to_color


def __find_text_in_runs(runs: List[Run],
                        text: str
                        ) -> Generator[Tuple[Run, str], None, None]:
    # FIXME красит лишнее если run заканчивается, пара букв в него попала,
    #  но в следующем run нет продолжения. Безумно редкий случай,
    #  скорее всего, можно создать только искусственно (см. template.docx)
    text_symbols = list(text)
    for run in runs:
        run_contains: List[str] = []
        for run_symbol in run.text:
            try:
                symbol = text_symbols.pop(Index.first)
                if run_symbol != symbol:
                    run_contains.clear()
                    text_symbols = __text_symbols_renew(text)
                else:
                    run_contains.append(symbol)
            except IndexError:
                if run_contains:
                    yield run, ''.join(run_contains)
                run_contains.clear()
                text_symbols = __text_symbols_renew(text)
                continue
        if run_contains:
            yield run, ''.join(run_contains)


def __text_symbols_renew(text: str) -> list[str]:
    return list(text)


def __reshape_run_with_text(paragraph: Paragraph, run: Run, text: str) -> Run:
    run_with_text_after_split_index = 1
    runs = paragraph.runs
    run_index = [r.text for r in runs].index(run.text)
    new_runs = __split_run(run, text)
    paragraph.clear()
    __add_runs(paragraph, runs[:run_index] + new_runs + runs[run_index + 1:])
    return new_runs[run_with_text_after_split_index]


def __split_run(run: Run, text: str) -> list[Run]:
    first_r = deepcopy(run)
    second_r = deepcopy(run)
    third_r = deepcopy(run)
    first_r.text, third_r.text = run.text.split(text, maxsplit=1)
    second_r.text = text
    return [first_r, second_r, third_r]


def __add_runs(paragraph: Paragraph, runs: list[Run]) -> None:
    runs_number = len(paragraph.runs)
    paragraph.append_runs(runs)
    # append_runs добавляет Run(' ') в начало, убираем следующей строкой
    paragraph.runs[runs_number].clear()


def __color_run(run: Run, color: str) -> None:
    run.font.color.rgb = Color.get(color)
